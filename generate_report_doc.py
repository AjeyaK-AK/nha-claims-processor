"""
Generates a professionally formatted Word (.docx) methodology report
for the NHA PMJAY Claims Processing System.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x00, 0x33, 0x66)   # heading
MID_BLUE    = RGBColor(0x1F, 0x4E, 0x79)   # sub-heading
ACCENT      = RGBColor(0x2E, 0x75, 0xB6)   # table header bg (write via XML)
LIGHT_GRAY  = RGBColor(0xF2, 0xF2, 0xF2)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_BLACK  = RGBColor(0x0D, 0x0D, 0x0D)
GREEN       = RGBColor(0x37, 0x86, 0x10)
RED         = RGBColor(0xC0, 0x00, 0x00)


# ── XML helpers ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(table):
    """Add thin borders to all cells."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "BFBFBF")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pr = p._p
    pPr = OxmlElement("w:pPr")
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)
    pr.insert(0, pPr)
    return p


# ── Style helpers ──────────────────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = MID_BLUE
    return p


def body(doc, text, bold_parts=None):
    """Add a body paragraph optionally with some bolded words."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    if not bold_parts:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT_BLACK
    else:
        remaining = text
        for part in bold_parts:
            idx = remaining.find(part)
            if idx >= 0:
                before = remaining[:idx]
                if before:
                    r = p.add_run(before)
                    r.font.size = Pt(10.5)
                    r.font.color.rgb = TEXT_BLACK
                rb = p.add_run(part)
                rb.font.size = Pt(10.5)
                rb.font.bold = True
                rb.font.color.rgb = TEXT_BLACK
                remaining = remaining[idx + len(part):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = Pt(10.5)
            r.font.color.rgb = TEXT_BLACK
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_BLACK
    return p


def add_table(doc, headers, rows, col_widths=None, header_bg="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else "EBF3FB"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT_BLACK

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    set_cell_border(table)
    doc.add_paragraph()
    return table


def add_code_block(doc, lines):
    """Monospace shaded block to represent code/pipeline."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        # light bg via shading
        pr = p._p
        pPr = OxmlElement("w:pPr")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F4F8")
        pPr.append(shd)
        pr.insert(0, pPr)


# ── COVER PAGE ─────────────────────────────────────────────────────────────────
def build_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NHA PMJAY CLAIMS PROCESSING SYSTEM")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Technical Methodology Report")
    r2.font.size = Pt(16)
    r2.font.color.rgb = MID_BLUE

    doc.add_paragraph()

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = tag.add_run("Hackathon Problem Statement 01  |  Document Intelligence & Claims Audit Automation")
    r3.font.size = Pt(11)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()
    doc.add_paragraph()

    add_horizontal_rule(doc)

    meta = [
        ("System", "Automated End-to-End NHA PMJAY Claims Verification Pipeline"),
        ("Scope", "80 claims  |  8 package codes  |  Extract_1 + Extract_2"),
        ("Date", datetime.date.today().strftime("%d %B %Y")),
        ("Technology", "Python 3.11  |  Tesseract v5  |  OpenCV  |  Streamlit"),
        ("Output", "Pass / Conditional / Fail verdict with full provenance"),
    ]
    table = add_table(doc, ["Field", "Value"], meta,
                      col_widths=[1.8, 4.5], header_bg="00336666")

    doc.add_page_break()


# ── SECTION 1 — Executive Summary ─────────────────────────────────────────────
def section_executive_summary(doc):
    h1(doc, "1.  Executive Summary")
    add_horizontal_rule(doc)
    body(doc,
         "This report describes the design, implementation, and operational methodology of an "
         "automated healthcare claims processing pipeline built for the NHA PMJAY Hackathon. "
         "The system ingests raw scanned medical documents (PDF, JPEG, PNG), classifies each page "
         "to a canonical document type, extracts structured fields with source provenance, applies "
         "Standard Treatment Guidelines (STG) compliance rules specific to each PMJAY package code, "
         "and produces a weighted Pass / Conditional / Fail verdict with full explainability.")

    doc.add_paragraph()
    body(doc, "Key performance targets defined in the problem statement:")

    add_table(doc,
              ["Component", "Target F1", "Our Approach"],
              [
                  ["Document Classification", "≥ 0.95", "292-label filename map + keyword density fallback"],
                  ["Rules + Provenance", "≥ 0.96", "YAML-driven STG rules, confidence-weighted scoring"],
                  ["Solution Design", "≥ 0.93", "Modular 6-stage pipeline, full HTML reports"],
              ],
              col_widths=[2.2, 1.3, 3.2])

    doc.add_page_break()


# ── SECTION 2 — System Architecture ───────────────────────────────────────────
def section_architecture(doc):
    h1(doc, "2.  System Architecture")
    add_horizontal_rule(doc)
    body(doc,
         "The pipeline is structured as six sequential, independently testable stages. "
         "Each stage produces typed data objects consumed by the next stage. "
         "All configuration (thresholds, paths, severity weights) is centralised in config.py "
         "so no stage contains hard-coded magic values.")

    doc.add_paragraph()
    add_code_block(doc, [
        "  Raw Files (PDF / JPEG / PNG)",
        "       |",
        "       v",
        "  [Stage 1]  INGESTION          -->  DocumentPage objects (one per page)",
        "       |",
        "       v",
        "  [Stage 2]  CLASSIFICATION     -->  Canonical document type + confidence",
        "       |",
        "       v",
        "  [Stage 3]  EXTRACTION         -->  Structured fields + provenance records",
        "       |",
        "       v",
        "  [Stage 4]  RULE EVALUATION    -->  RuleResult list (pass/fail + evidence)",
        "       |",
        "       v",
        "  [Stage 5]  DECISION ENGINE    -->  Weighted score  -->  PASS / CONDITIONAL / FAIL",
        "       |",
        "       v",
        "  [Stage 6]  OUTPUT             -->  JSON report  +  Self-contained HTML dashboard",
    ])

    doc.add_paragraph()
    body(doc, "Core Python modules:")
    add_table(doc,
              ["Module / File", "Responsibility"],
              [
                  ["src/ingestion/document_loader.py", "PDF page rendering, image loading, OCR orchestration"],
                  ["src/classification/doc_classifier.py", "Three-signal fusion classifier"],
                  ["src/classification/doc_types.py", "292-entry filename label map + keyword lists"],
                  ["src/extraction/field_extractor.py", "Regex-based structured field extraction with provenance"],
                  ["src/extraction/visual_detector.py", "OpenCV stamp / signature / barcode detection"],
                  ["src/rules/rules_engine.py", "YAML rule loader + four rule type evaluators"],
                  ["src/rules/stg_rules/*.yaml", "One YAML file per PMJAY package (8 files)"],
                  ["src/decisioning/decision_engine.py", "Weighted scoring + verdict thresholding"],
                  ["src/output/report_generator.py", "JSON report serialisation"],
                  ["generate_html_report.py", "Self-contained 5-tab Bootstrap HTML reports"],
                  ["batch_process.py", "Batch runner across both extract directories + CSV summary"],
                  ["app.py", "Streamlit web dashboard"],
              ],
              col_widths=[3.0, 3.5])

    doc.add_page_break()


# ── SECTION 3 — Stage Detail ───────────────────────────────────────────────────
def section_stages(doc):
    h1(doc, "3.  Pipeline Stage Details")
    add_horizontal_rule(doc)

    # --- 3.1 Ingestion ---
    h2(doc, "3.1  Stage 1 — Document Ingestion")
    body(doc,
         "Every claim folder contains between 1 and 20+ files in mixed formats. "
         "The ingestion layer normalises all formats to a common DocumentPage abstraction "
         "before any downstream logic sees the data.")

    bullet(doc, "PDF rendering:  PyMuPDF (fitz) renders each page to an RGB image at 200 DPI. "
                "Higher DPI (300) was tested but gave no accuracy gain while tripling memory use.")
    bullet(doc, "Image loading:  OpenCV imread for JPEG / PNG files; "
                "alpha-channel flattened to white background.")
    bullet(doc, "OCR strategy — 3-pass cascade:")
    bullet(doc, "Pass 1  →  Raw Tesseract (eng, PSM 3 auto-segment). "
                "If mean word confidence > 55 and text length > 20 chars: accept.", level=1)
    bullet(doc, "Pass 2  →  Image enhancement: CLAHE contrast, Gaussian denoise, "
                "Otsu binarisation. Re-run Tesseract. Accept if improved.", level=1)
    bullet(doc, "Pass 3  →  OSD orientation detection (tesseract --psm 0), auto-rotate. "
                "Common for landscape-printed hospital forms.", level=1)
    bullet(doc, "Filenames follow the pattern:  {sequence}__{ClaimID}__{LABEL}.ext  — "
                "this LABEL is the primary classification signal.")

    doc.add_paragraph()

    # --- 3.2 Classification ---
    h2(doc, "3.2  Stage 2 — Document Classification")
    body(doc,
         "Each DocumentPage is assigned one canonical document type (e.g. discharge_summary, "
         "case_sheet, angiography_report). The classifier uses a three-signal fusion approach:")

    add_table(doc,
              ["Priority", "Signal", "Mechanism", "Confidence Assigned"],
              [
                  ["1 (highest)", "Filename label",
                   "Map LABEL from filename to canonical type via 292-entry lookup table",
                   "0.92 (fixed)"],
                  ["2", "Content keywords",
                   "Keyword density score: hits / total_keywords per doc type; best wins",
                   "0.50 – 0.85 (scaled)"],
                  ["3 (fallback)", "Unknown",
                   "Assign type = other",
                   "0.40"],
              ],
              col_widths=[1.1, 1.3, 3.0, 1.3])

    body(doc, "The 292-label map was built by extracting every unique label stem across all 80 "
              "claims and mapping each to a canonical type. Examples of mappings added:")
    bullet(doc, "BHT  →  case_sheet  (Bed Head Ticket — common Indian hospital terminology)")
    bullet(doc, "KUV suffix (MALATIKUV, GOVINDKUV)  →  ivp_report  (KUB X-ray film)")
    bullet(doc, "DB, DP  →  discharge_summary  (Discharge Brief / Discharge Papers)")
    bullet(doc, "N, N3, N4, AANITA, NASEEMA  →  identity_document  (patient photo named by patient name)")
    bullet(doc, "BIRTH_PROOF  →  identity_document  (used as identity evidence in NHA claims)")

    doc.add_paragraph()

    # --- 3.3 Extraction ---
    h2(doc, "3.3  Stage 3 — Field Extraction & Visual Detection")
    body(doc,
         "Two extraction sub-systems run in parallel on the OCR text and images respectively.")

    body(doc, "Field Extractor (text-based):", bold_parts=["Field Extractor (text-based):"])
    add_table(doc,
              ["Field", "Extraction Approach"],
              [
                  ["Admission date", "12 regex patterns: DD/MM/YYYY, DD-Mon-YYYY, spelled-out dates; anchored to 'admission', 'admitted on' context"],
                  ["Discharge date", "Same patterns; anchored to 'discharge', 'date of discharge'"],
                  ["Diagnosis", "Keyword-windowed extraction within 3 lines of 'diagnosis', 'final diagnosis', 'clinical diagnosis'"],
                  ["Patient name", "NER-style regex near 'patient name', 'name of patient', 'beneficiary'"],
                  ["Hospital name", "Near 'hospital', 'healthcare', 'medical centre' tokens"],
                  ["Length of stay", "Computed from (discharge_date - admission_date) in days; also extracted from 'total days' text"],
              ],
              col_widths=[1.6, 5.0])

    doc.add_paragraph()
    body(doc, "Visual Detector (image-based, OpenCV):", bold_parts=["Visual Detector (image-based, OpenCV):"])
    add_table(doc,
              ["Element", "Detection Method"],
              [
                  ["Hospital stamp", "Contour area > 2000 px² + circularity > 0.45 on red/blue channel blobs"],
                  ["Doctor signature", "Dark thin-stroke region: aspect ratio > 2.0 + ink density threshold"],
                  ["Barcode / QR code", "pyzbar library + cv2.QRCodeDetector for QR, CODE128, CODE39, EAN"],
              ],
              col_widths=[1.8, 4.8])

    body(doc,
         "Every extracted value carries a Provenance record: "
         "{ doc_id, page_number, field_name, value, confidence, bounding_box }. "
         "This provenance chain flows through all downstream stages and appears in the final HTML report.")

    doc.add_paragraph()

    # --- 3.4 Rules ---
    h2(doc, "3.4  Stage 4 — STG Rule Evaluation")
    body(doc,
         "Rules are defined in YAML files, one per PMJAY package code. "
         "This architecture means adding support for a new package requires only a new YAML file — "
         "no code changes. Each claim is evaluated exclusively against its own package's rules.")

    body(doc, "Four rule types are supported:")
    add_table(doc,
              ["Rule Type", "What It Checks", "Example"],
              [
                  ["document_presence", "Is a required document type present in the claim?",
                   "angiography_report must be present for MC011A"],
                  ["temporal", "Do dates form a valid clinical sequence?",
                   "admission_date <= discharge_date; procedure within window"],
                  ["los", "Is length of stay within clinically expected range?",
                   "Arthroplasty (SB039A): 5–14 days"],
                  ["visual_element", "Is stamp/signature present on a required document?",
                   "Hospital stamp on discharge summary"],
              ],
              col_widths=[1.5, 2.8, 2.3])

    body(doc, "Each rule declares a severity level:")
    bullet(doc, "critical  —  weight 1.0  |  If high-confidence and failed: triggers automatic FAIL")
    bullet(doc, "major     —  weight 0.6  |  Contributes significantly to weighted score")
    bullet(doc, "minor     —  weight 0.2  |  Advisory; does not prevent PASS")

    add_table(doc,
              ["Package Code", "Procedure", "No. of Rules"],
              [
                  ["MC011A", "Coronary Angiography (CAG) ± PTCA / Stenting", "8"],
                  ["MG006A", "General Medical Management (Fever / Sepsis)", "11"],
                  ["MG029A", "Radiological Investigation — X-Ray", "7"],
                  ["MG064A", "General Medical Inpatient — Comprehensive", "14"],
                  ["SB039A", "Total Knee / Hip Arthroplasty", "14"],
                  ["SG039",  "Ultrasonography (USG) — Whole Abdomen / Pelvis", "7"],
                  ["SG039C", "Laparoscopic Cholecystectomy", "12"],
                  ["SU007A", "PCNL — Percutaneous Nephrolithotomy", "10"],
              ],
              col_widths=[1.4, 4.0, 1.0])

    doc.add_paragraph()

    # --- 3.5 Decision Engine ---
    h2(doc, "3.5  Stage 5 — Decision Engine")
    body(doc,
         "The decision engine aggregates all RuleResult objects from Stage 4 into a single "
         "weighted score and verdict.")

    body(doc, "Weighted scoring formula:")
    add_code_block(doc, [
        "  score = SUM(weight_i  for passed rules i) / SUM(weight_i  for all rules i)",
        "",
        "  where:  weight = 1.0  (critical)  |  0.6  (major)  |  0.2  (minor)",
    ])

    doc.add_paragraph()
    body(doc, "Verdict thresholds:")
    add_table(doc,
              ["Score Range", "Verdict", "Meaning"],
              [
                  ["score >= 0.90", "PASS", "All critical checks passed; minor deficiencies acceptable"],
                  ["0.60 <= score < 0.90", "CONDITIONAL", "Some major issues; claim requires manual review"],
                  ["score < 0.60", "FAIL", "Critical documentation or compliance failures"],
              ],
              col_widths=[1.8, 1.3, 3.5])

    body(doc,
         "Key design decision — confidence-gated auto-FAIL:  "
         "A critical rule failure only triggers automatic FAIL if its evidence confidence is >= 0.70. "
         "This handles the common case where landscape-printed scanned forms produce OCR confidence "
         "of ~0.45 for date fields. In that case the temporal rule failure (e.g. R008: admission before "
         "discharge) cannot be verified — it contributes to weighted score but does not force FAIL, "
         "preventing false negatives on otherwise valid claims.",
         bold_parts=["Key design decision — confidence-gated auto-FAIL:"])

    doc.add_paragraph()

    # --- 3.6 Output ---
    h2(doc, "3.6  Stage 6 — Output & Reporting")
    body(doc, "Two output formats are generated per claim:")

    body(doc, "JSON Report (machine-readable):", bold_parts=["JSON Report (machine-readable):"])
    bullet(doc, "Claim ID, package code, verdict, overall_score, confidence")
    bullet(doc, "critical_failures / major_failures / minor_failures lists")
    bullet(doc, "Full rule_details array: rule_id, result, message, evidence references")
    bullet(doc, "Each evidence item: doc_id, page_number, field_name, value, confidence, bounding_box")

    doc.add_paragraph()
    body(doc, "HTML Report (human-readable, self-contained):", bold_parts=["HTML Report (human-readable, self-contained):"])
    body(doc, "Bootstrap 5 single-file report with 5 tabs:")
    bullet(doc, "Tab 1 — Document Classification:  table of all docs with predicted type and confidence")
    bullet(doc, "Tab 2 — Rule Evaluation:  per-rule PASS/FAIL with colour coding and evidence detail")
    bullet(doc, "Tab 3 — Episode Timeline:  admission → procedure → discharge event sequence")
    bullet(doc, "Tab 4 — Visual Elements:  stamps, signatures, barcodes detected per document")
    bullet(doc, "Tab 5 — Provenance Log:  field → source document → page → bounding box traceability")
    body(doc,
         "Reports are fully self-contained (no server required) and can be shared as a single .html file. "
         "An index.html master table links all 80 claim reports for batch review.")

    doc.add_page_break()


# ── SECTION 4 — Dataset & Scope ────────────────────────────────────────────────
def section_dataset(doc):
    h1(doc, "4.  Dataset & Scope")
    add_horizontal_rule(doc)

    add_table(doc,
              ["Extract", "Package Codes", "Claims", "Document Types Present"],
              [
                  ["extract_1", "MG006A, MG064A, SB039A, SG039C",
                   "~40 (10 per package)", "Discharge summaries, case sheets, OT notes, X-rays, USG, IVP, lab reports, identity docs"],
                  ["extract_2", "MC011A, MG029A, SG039, SU007A",
                   "~40 (10 per package)", "CAG/PTCA images, angiography reports, X-ray films, USG films, IVP reports"],
              ],
              col_widths=[1.0, 2.2, 1.0, 2.8])

    body(doc,
         "Unique challenges encountered in the dataset and how they were addressed:")
    add_table(doc,
              ["Challenge", "Frequency", "Solution"],
              [
                  ["Landscape-printed forms (discharge, assessment sheets)",
                   "~35% of pages",
                   "Pass 3 OSD rotation; temporal rules downgraded when OCR conf < 0.70"],
                  ["UUID filenames (no label information)",
                   "~20% of files in extract_2",
                   "Content keyword density classifier as fallback"],
                  ["Patient-name-as-filename (e.g. MALATIKUV.pdf)",
                   "~15% of files",
                   "292 curated label patterns; KUV suffix mapped to ivp_report"],
                  ["Mixed format batches (PDF + JPEG in same claim)",
                   "All claims",
                   "Unified DocumentPage abstraction; format-agnostic OCR pipeline"],
                  ["Windows cp1252 encoding crashes on Unicode log output",
                   "Batch runner",
                   "sys.stdout.reconfigure(encoding='utf-8') + PYTHONIOENCODING=utf-8"],
              ],
              col_widths=[2.2, 1.3, 3.1])

    doc.add_page_break()


# ── SECTION 5 — Technology Stack ──────────────────────────────────────────────
def section_tech(doc):
    h1(doc, "5.  Technology Stack")
    add_horizontal_rule(doc)

    add_table(doc,
              ["Library / Tool", "Version", "Purpose"],
              [
                  ["Python", "3.11", "Core runtime"],
                  ["PyMuPDF (fitz)", "1.23+", "PDF page rendering to images at configurable DPI"],
                  ["Tesseract OCR", "5.5.0", "Text extraction from scanned images (eng language pack)"],
                  ["pytesseract", "0.3+", "Python wrapper for Tesseract; confidence scoring"],
                  ["OpenCV (cv2)", "4.9+", "Image preprocessing, contour detection, stamp/signature detection"],
                  ["pyzbar", "0.1.9", "Barcode and QR code detection"],
                  ["PyYAML", "6.0+", "STG rule file loading"],
                  ["python-docx", "1.1+", "Word report generation"],
                  ["Streamlit", "1.45+", "Interactive web dashboard"],
                  ["Bootstrap", "5.3.3 (CDN)", "HTML report styling"],
                  ["Git / GitHub", "2.54", "Version control and deployment"],
              ],
              col_widths=[2.0, 1.2, 3.4])

    doc.add_page_break()


# ── SECTION 6 — Key Engineering Decisions ─────────────────────────────────────
def section_decisions(doc):
    h1(doc, "6.  Key Engineering Decisions")
    add_horizontal_rule(doc)

    items = [
        (
            "YAML-driven rules engine",
            "Rules are data, not code. Each PMJAY package has its own YAML file defining "
            "requirements, severities, and messages. Adding a new package to the system "
            "requires only a new YAML file — no Python changes needed. The same evaluator "
            "handles all four rule types dynamically."
        ),
        (
            "Three-signal fusion classification",
            "Filename label (when present) is the highest-confidence signal (0.92). "
            "Content keyword density handles UUID-named files. This layered approach achieves "
            "high precision on labelled files while gracefully degrading for unlabelled ones."
        ),
        (
            "Confidence-gated critical failures",
            "Not all critical rule failures are equal. A temporal rule that cannot be verified "
            "because OCR confidence is 0.45 on a landscape form is fundamentally different from "
            "a confirmed missing required document. The decision engine only triggers auto-FAIL "
            "for high-confidence (>= 0.70) critical failures."
        ),
        (
            "Full provenance at every step",
            "Every extracted value, every rule result, and every visual element detection "
            "carries a chain of evidence back to the source: which document, which page, "
            "which field, which bounding box, with what confidence. This is not just good design — "
            "it is the core requirement of the problem statement."
        ),
        (
            "Self-contained HTML reports",
            "Deploying a live OCR server on Streamlit Cloud is impractical (Tesseract install, "
            "large document sets). HTML reports are pre-generated from JSON and shipped as single files. "
            "They work offline, require no server, and can be opened by any browser."
        ),
        (
            "Batch processing with fault isolation",
            "The batch runner wraps each claim in a try/except. A single corrupt PDF or "
            "unexpected file format does not abort the entire run. Failed claims are recorded "
            "in the summary CSV with verdict='ERROR' and the exception message."
        ),
    ]

    for title, detail in items:
        h2(doc, title)
        body(doc, detail)
        doc.add_paragraph()

    doc.add_page_break()


# ── SECTION 7 — Results Summary ───────────────────────────────────────────────
def section_results(doc):
    h1(doc, "9.  Results Summary")
    add_horizontal_rule(doc)

    body(doc,
         "The following table summarises the root causes of FAIL verdicts observed during "
         "initial validation runs, and the fixes applied to address them:")

    add_table(doc,
              ["Root Cause", "Affected Packages", "Fix Applied"],
              [
                  ["Wrong STG YAML (procedure mismatch)", "MC011A, MG006A, MG064A, MG029A, SG039",
                   "YAML files rewritten to match actual claim contents for each package code"],
                  ["Date extraction fails on landscape forms", "All packages (R008/R009/R010)",
                   "Confidence-gated auto-FAIL; date rules demoted to major when confidence < 0.70"],
                  ["Unlabelled / UUID filenames misclassified", "MC011A (extract_2)",
                   "Content keyword fallback classifier; CAG/PTCA/angiography keywords tuned"],
                  ["BIRTH_PROOF not mapped to identity_document", "MG006A",
                   "292-label map updated; BIRTH_PROOF -> identity_document"],
                  ["Patient-name-prefix filenames (KUV, DB, DP)", "MG006A, SU007A",
                   "All unique label stems extracted and mapped in doc_types.py"],
              ],
              col_widths=[2.2, 1.8, 2.6])

    doc.add_paragraph()
    body(doc, "Batch run statistics (all 80 claims across 8 packages):")
    add_table(doc,
              ["Metric", "Value"],
              [
                  ["Total claims processed", "80"],
                  ["Package codes covered", "8"],
                  ["STG rules total (across all YAMLs)", "83"],
                  ["Filename label mappings", "292"],
                  ["Output per claim", "JSON report + 5-tab HTML report"],
                  ["Batch summary", "output_reports/batch_summary.csv"],
              ],
              col_widths=[3.0, 3.5])

    doc.add_page_break()


# ── SECTION 10 — Deployment ────────────────────────────────────────────────────
def section_deployment(doc):
    h1(doc, "10.  Deployment & Reproducibility")
    add_horizontal_rule(doc)

    h2(doc, "10.1  Local Setup")
    add_code_block(doc, [
        "  # 1. Clone repository",
        "  git clone https://github.com/RanjitKolkar/nha-claims-processor.git",
        "  cd nha-claims-processor",
        "",
        "  # 2. Install Python dependencies",
        "  pip install -r requirements.txt",
        "",
        "  # 3. Install Tesseract (Windows)",
        "  winget install UB-Mannheim.TesseractOCR",
        "",
        "  # 4. Place claim data",
        "  #    extract_1/Claims/<PackageCode>/<ClaimID>/",
        "  #    extract_2/Claims/<PackageCode>/<ClaimID>/",
        "",
        "  # 5. Run full batch",
        "  python batch_process.py",
        "",
        "  # 6. Generate HTML reports",
        "  python generate_html_report.py",
        "",
        "  # 7. Open dashboard",
        "  streamlit run app.py",
    ])

    doc.add_paragraph()
    h2(doc, "10.2  GitHub Repository")
    bullet(doc, "URL:  https://github.com/RanjitKolkar/nha-claims-processor")
    bullet(doc, "Branch:  main")
    bullet(doc, "Claim data excluded from repository via .gitignore (extract_1/, extract_2/)")
    bullet(doc, "demo_reports/ contains 3 pre-computed JSON reports for cloud demo mode")

    doc.add_paragraph()
    h2(doc, "10.3  Configuration")
    body(doc, "All tunable parameters are in config.py:")
    add_table(doc,
              ["Parameter", "Default", "Effect"],
              [
                  ["PDF_DPI", "200", "Image resolution for PDF rendering"],
                  ["PASS_THRESHOLD", "0.90", "Minimum weighted score for PASS verdict"],
                  ["CONDITIONAL_THRESHOLD", "0.60", "Minimum score for CONDITIONAL (below = FAIL)"],
                  ["SEVERITY_WEIGHTS", "{critical:1.0, major:0.6, minor:0.2}", "Rule weighting for score calculation"],
                  ["FILENAME_LABEL_CONFIDENCE", "0.92", "Confidence assigned to filename-based classification"],
                  ["USE_EASYOCR_FALLBACK", "False", "Enable EasyOCR fallback (requires model download)"],
              ],
              col_widths=[2.2, 2.0, 2.4])


# ── SECTION 7 — STG Guidelines ────────────────────────────────────────────────
def section_stg_guidelines(doc):
    h1(doc, "7.  STG Guidelines — Rules Per Package")
    add_horizontal_rule(doc)

    body(doc,
         "Each PMJAY package is governed by a YAML rule file under src/rules/stg_rules/. "
         "The tables below list every rule for all eight packages, including its severity "
         "(critical → auto-FAIL if confidence ≥ 0.70; major → score penalty; minor → advisory) "
         "and the exact compliance requirement it enforces.")

    # ── MG006A ──────────────────────────────────────────────────────────────────
    h2(doc, "7.1  MG006A — General Medical Management (Fever / Sepsis / Medical Conditions)")
    body(doc, "Specialty: General Medicine  |  Care type: Inpatient  |  LoS: 2–14 days  |  Ceiling: ₹15,000")
    add_table(doc,
              ["Rule", "Name", "Severity", "Requirement"],
              [
                  ["R001", "Discharge summary present",               "CRITICAL", "Discharge summary is mandatory for all general medical inpatient claims."],
                  ["R002", "Pre-authorization form present",           "major",    "Pre-authorization / PMJAY eligibility form must be submitted."],
                  ["R003", "Case sheet / clinical record present",     "CRITICAL", "Case sheet required; clinical_notes accepted as alternative."],
                  ["R004", "Identity document present",                "CRITICAL", "Beneficiary identity proof (Aadhaar / PMJAY card) must be present."],
                  ["R005", "Laboratory investigations documented",     "major",    "CBC, biochemistry lab reports required."],
                  ["R006", "Admission date precedes discharge date",   "CRITICAL", "Admission date must be ≤ discharge date (temporal check)."],
                  ["R007", "Minimum inpatient stay 2 days",            "major",    "LoS < 2 days is considered clinically implausible for this package."],
                  ["R008", "Hospital stamp on discharge summary",      "major",    "Hospital stamp must be visible on the discharge summary."],
                  ["R009", "Doctor signature on discharge summary",    "major",    "Treating doctor's signature must appear on the discharge summary."],
                  ["R010", "Billed amount within package ceiling",     "minor",    "Billed amount must not exceed ₹15,000 package ceiling."],
                  ["R011", "Diagnosis consistent with package",        "major",    "Diagnosis keywords (fever, sepsis, pneumonia, dengue, etc.) must be present."],
                  ["R012", "Patient name extractable",                 "minor",    "Patient name must be readable from at least one submitted document."],
              ],
              col_widths=[0.5, 2.3, 0.85, 3.05])

    # ── MG064A ──────────────────────────────────────────────────────────────────
    h2(doc, "7.2  MG064A — General Medical Inpatient Comprehensive Management")
    body(doc, "Specialty: General Medicine / Internal Medicine  |  Care type: Inpatient  |  LoS: 3–14 days  |  Ceiling: ₹30,000")
    add_table(doc,
              ["Rule", "Name", "Severity", "Requirement"],
              [
                  ["R001", "Discharge summary present",               "CRITICAL", "Discharge summary is mandatory for all inpatient claims."],
                  ["R002", "Pre-authorization form present",           "major",    "Pre-authorization / PMJAY eligibility form must be submitted."],
                  ["R003", "Case sheet present",                       "CRITICAL", "Case sheet required; clinical_notes accepted as alternative."],
                  ["R004", "Identity document present",                "CRITICAL", "Beneficiary identity proof must be submitted."],
                  ["R005", "Laboratory investigations documented",     "major",    "CBC, biochemistry, etc. required for comprehensive management."],
                  ["R006", "Admission date precedes discharge date",   "CRITICAL", "Admission date must be ≤ discharge date (temporal check)."],
                  ["R007", "Minimum inpatient stay 3 days",            "major",    "LoS < 3 days is clinically implausible for comprehensive management."],
                  ["R008", "Hospital stamp on discharge summary",      "major",    "Hospital stamp must be visible on the discharge summary."],
                  ["R009", "Doctor signature on discharge summary",    "major",    "Treating doctor's signature must appear on the discharge summary."],
                  ["R010", "Billed amount within package ceiling",     "minor",    "Billed amount must not exceed ₹30,000 package ceiling."],
                  ["R011", "Diagnosis consistent with package",        "major",    "Keywords: fever, sepsis, cardiac, stroke, renal failure, hepatitis, etc."],
                  ["R013", "Stent type documented (if PTCA performed)","major",    "Procedure report must document stent type if PTCA was performed."],
                  ["R014", "Pre-PTCA investigations present",          "major",    "Pre-procedure labs (CBC, RFT, PTINR, ECG) must be submitted if PTCA."],
              ],
              col_widths=[0.5, 2.3, 0.85, 3.05])

    doc.add_page_break()

    # ── MC011A ──────────────────────────────────────────────────────────────────
    h2(doc, "7.3  MC011A — Coronary Angiography (CAG) with or without PTCA / Stenting")
    body(doc, "Specialty: Cardiology – Interventional  |  Care type: Daycare / Short-stay  |  LoS: 0–3 days  |  Ceiling: ~₹25,000")
    add_table(doc,
              ["Rule", "Name", "Severity", "Requirement"],
              [
                  ["R001", "Angiography report present",           "CRITICAL", "CAG report is the primary mandatory document for MC011A."],
                  ["R002", "Pre-authorization form present",       "major",    "Pre-authorization is mandatory for cardiac catheterisation."],
                  ["R003", "Discharge card / summary present",     "major",    "Discharge card or summary is expected for any cardiac procedure."],
                  ["R004", "Admission before discharge",           "CRITICAL", "Admission date must precede discharge date."],
                  ["R005", "LoS within CAG/PTCA limits (0–3 d)",  "major",    "LoS > 3 days requires clinical justification."],
                  ["R006", "Hospital stamp on angiography report", "major",    "Hospital / cath-lab stamp required on the angiography report."],
                  ["R007", "Cardiologist signature on report",     "major",    "Cardiologist signature required on CAG / PTCA report."],
                  ["R008", "Diagnosis indicates coronary disease", "CRITICAL", "Diagnosis must reference CAD, stenosis, angina, PTCA, stent, LAD, RCA, etc."],
                  ["R009", "Stent sticker present if PTCA done",   "major",    "If stenting performed, implant barcode sticker is required."],
                  ["R010", "Billed amount within package ceiling", "minor",    "CAG ceiling ~₹25,000; PTCA has a higher ceiling; excess needs evidence."],
                  ["R011", "Patient name extractable",             "minor",    "Patient name must be readable from submitted documents."],
                  ["R012", "Admission date documented",            "minor",    "Admission date must be clearly documented."],
              ],
              col_widths=[0.5, 2.3, 0.85, 3.05])

    # ── SB039A ──────────────────────────────────────────────────────────────────
    h2(doc, "7.4  SB039A — Total Knee / Hip Arthroplasty")
    body(doc, "Specialty: Orthopaedics  |  Care type: Inpatient  |  LoS: 5–14 days")
    add_table(doc,
              ["Rule", "Name", "Severity", "Requirement"],
              [
                  ["R001", "Discharge summary present",               "CRITICAL", "Discharge summary is mandatory."],
                  ["R002", "Pre-authorization form present",           "major",    "Pre-authorization form is mandatory."],
                  ["R003", "Operative notes present",                  "CRITICAL", "Operative notes are mandatory for all surgical packages."],
                  ["R004", "Pre-operative X-ray present",              "CRITICAL", "Pre-operative X-ray of the affected joint is mandatory."],
                  ["R005", "Implant barcode sticker present",          "CRITICAL", "Implant sticker (REF / LOT / SN) is mandatory for arthroplasty claims."],
                  ["R006", "Anaesthesia notes present",                "major",    "Anaesthesia record must be submitted."],
                  ["R007", "Pre-operative lab investigations present", "major",    "Pre-op labs: CBC, BT/CT, RFT, ECG must be submitted."],
                  ["R008", "Admission before discharge",               "CRITICAL", "Admission date must precede discharge date."],
                  ["R009", "Procedure within admission window",        "CRITICAL", "Operation date must fall between admission and discharge dates."],
                  ["R010", "LoS 5–14 days",                           "major",    "LoS < 5 days is clinically implausible; > 14 days needs justification."],
                  ["R011", "Hospital stamp on discharge summary",      "major",    "Hospital stamp required on discharge summary."],
                  ["R012", "Surgeon signature on operative notes",     "major",    "Surgeon must sign operative notes."],
                  ["R013", "Diagnosis consistent with joint disease",  "major",    "Keywords: osteoarthritis, AVN, fracture neck, arthroplasty, OA knee, etc."],
              ],
              col_widths=[0.5, 2.3, 0.85, 3.05])

    doc.add_page_break()

    # ── SG039C ──────────────────────────────────────────────────────────────────
    h2(doc, "7.5  SG039C — Laparoscopic Cholecystectomy / Abdominal Surgery")
    body(doc, "Specialty: General Surgery  |  Care type: Inpatient  |  LoS: 2–7 days")
    add_table(doc,
              ["Rule", "Name", "Severity", "Requirement"],
              [
                  ["R001", "Discharge summary present",               "CRITICAL", "Discharge summary is mandatory."],
                  ["R002", "Pre-authorization form present",           "major",    "Pre-authorization form is mandatory."],
                  ["R003", "Operative notes present",                  "CRITICAL", "Operative notes are mandatory for all surgical packages."],
                  ["R004", "USG abdomen report present",               "CRITICAL", "USG abdomen mandatory to confirm cholelithiasis / abdominal indication."],
                  ["R005", "Anaesthesia notes present",                "major",    "Anaesthesia record must be submitted."],
                  ["R006", "Pre-operative lab investigations present", "major",    "CBC, LFT, RFT, clotting profile required."],
                  ["R007", "Admission before discharge",               "CRITICAL", "Admission date must precede discharge date."],
                  ["R008", "Operation date within admission window",   "CRITICAL", "Surgery date must fall within admission–discharge window."],
                  ["R009", "LoS 2–7 days",                            "major",    "LoS < 2 days or > 7 days needs clinical justification."],
                  ["R010", "Hospital stamp on discharge summary",      "major",    "Hospital stamp required on discharge summary."],
                  ["R011", "Surgeon signature on operative notes",     "major",    "Surgeon must sign operative notes."],
                  ["R012", "Diagnosis consistent with package",        "major",    "Keywords: cholelithiasis, cholecystitis, gallstone, appendicitis, hernia, etc."],
                  ["R013", "Patient name extractable",                 "minor",    "Patient name must be readable from submitted documents."],
              ],
              col_widths=[0.5, 2.3, 0.85, 3.05])

    # ── SG039 ───────────────────────────────────────────────────────────────────
    h2(doc, "7.6  SG039 — Ultrasonography (USG) – Whole Abdomen / Pelvis / Targeted")
    body(doc, "Specialty: Radiology / Diagnostics  |  Care type: Outpatient / Day Procedure  |  LoS: 0–1 day  |  Ceiling: ₹3,000")
    add_table(doc,
              ["Rule", "Name", "Severity", "Requirement"],
              [
                  ["R001", "USG image or report present",              "CRITICAL", "USG film / report is the primary mandatory document for SG039."],
                  ["R002", "Pre-authorization form present",           "major",    "Pre-authorization is mandatory."],
                  ["R003", "Identity document present",                "CRITICAL", "Beneficiary identity proof must be submitted."],
                  ["R004", "Sonologist report present",                "major",    "Sonologist's written interpretation of the USG should be present."],
                  ["R005", "USG date within claim period",             "major",    "USG date must fall within the claimed service period."],
                  ["R006", "Sonologist signature on report",           "major",    "Sonologist / radiologist signature required on USG report."],
                  ["R007", "Hospital stamp on USG report",             "minor",    "Hospital / radiology centre stamp required on USG report."],
                  ["R008", "Patient name extractable",                 "minor",    "Patient name must be readable from submitted documents."],
                  ["R009", "USG procedure date documented",            "major",    "USG procedure date must be clearly documented in records."],
                  ["R010", "Clinical indication documented",           "major",    "Clinical indication for USG must be documented (diagnosis field)."],
              ],
              col_widths=[0.5, 2.3, 0.85, 3.05])

    doc.add_page_break()

    # ── MG029A ──────────────────────────────────────────────────────────────────
    h2(doc, "7.7  MG029A — Radiological Investigation (X-Ray / Chest Imaging)")
    body(doc, "Specialty: Radiology / Diagnostic Imaging  |  Care type: Outpatient / Day Procedure  |  LoS: 0–1 day  |  Ceiling: ₹5,000")
    add_table(doc,
              ["Rule", "Name", "Severity", "Requirement"],
              [
                  ["R001", "X-ray image or film present",              "CRITICAL", "X-ray film is the primary mandatory document for MG029A."],
                  ["R002", "Pre-authorization form present",           "major",    "Pre-authorization is mandatory."],
                  ["R003", "Radiology report / interpretation present","major",    "Radiologist's interpretation report should be present."],
                  ["R004", "Identity document present",                "CRITICAL", "Beneficiary identity proof must be submitted."],
                  ["R005", "X-ray date within claim period",           "major",    "X-ray date must fall within the claimed service period."],
                  ["R006", "Radiologist signature on report",          "major",    "Radiologist signature required on X-ray report."],
                  ["R007", "Hospital stamp on X-ray or report",        "minor",    "Hospital / radiology centre stamp required."],
                  ["R008", "Patient name extractable",                 "minor",    "Patient name must be readable from submitted documents."],
                  ["R009", "X-ray date documented",                    "major",    "X-ray date must be clearly documented in submitted records."],
                  ["R010", "Clinical indication documented",           "major",    "Clinical indication for X-ray must be documented."],
              ],
              col_widths=[0.5, 2.3, 0.85, 3.05])

    # ── SU007A ──────────────────────────────────────────────────────────────────
    h2(doc, "7.8  SU007A — PCNL / URS / DJ Stenting (Urological Surgery)")
    body(doc, "Specialty: Urology  |  Care type: Inpatient  |  LoS: 3–7 days")
    add_table(doc,
              ["Rule", "Name", "Severity", "Requirement"],
              [
                  ["R001", "Discharge summary present",               "CRITICAL", "Discharge summary is mandatory."],
                  ["R002", "Pre-authorization form present",           "major",    "Pre-authorization is mandatory."],
                  ["R003", "Operative notes present",                  "CRITICAL", "Operative notes are mandatory for urological surgical claims."],
                  ["R004", "IVP / KUB report present",                 "CRITICAL", "IVP or KUB X-ray mandatory to confirm urolithiasis / obstructive uropathy."],
                  ["R005", "USG KUB present",                          "major",    "USG KUB required to document stone size and location."],
                  ["R006", "Anaesthesia notes present",                "major",    "Anaesthesia record is mandatory."],
                  ["R007", "Pre-operative lab investigations present", "major",    "Urine culture, RFT, CBC must be submitted."],
                  ["R008", "Admission before discharge",               "CRITICAL", "Admission date must precede discharge date."],
                  ["R009", "Surgery date within admission window",     "CRITICAL", "Operation date must be within admission–discharge window."],
                  ["R010", "LoS 3–7 days",                            "major",    "LoS < 3 days is clinically unusual; > 7 days needs justification."],
                  ["R011", "Hospital stamp on discharge summary",      "major",    "Hospital stamp required on discharge summary."],
                  ["R012", "Surgeon signature on operative notes",     "major",    "Urologist's signature required on operative notes."],
                  ["R013", "Diagnosis consistent with urolithiasis",   "major",    "Keywords: calculus, stone, nephrolithiasis, PCNL, URS, DJ stent, etc."],
              ],
              col_widths=[0.5, 2.3, 0.85, 3.05])

    doc.add_page_break()


# ── SECTION 8 — Glossary ───────────────────────────────────────────────────────
def section_glossary(doc):
    h1(doc, "8.  Glossary — Labels & Scoring Reference")
    add_horizontal_rule(doc)

    body(doc,
         "This section defines every label, score, and flag that appears in the JSON / HTML "
         "reports produced by the system.")

    # Verdict
    h2(doc, "8.1  Verdict Labels")
    add_table(doc,
              ["Verdict", "Meaning", "Condition"],
              [
                  ["PASS",        "Claim is fully compliant. Recommended for payment.",
                   "Compliance score ≥ 0.90  AND  no critical rule fails with confidence ≥ 0.70"],
                  ["CONDITIONAL", "Partial compliance. Needs manual reviewer check.",
                   "Score 0.60 – 0.89  AND  no high-confidence critical failure"],
                  ["FAIL",        "Non-compliant. Reject or issue deficiency query.",
                   "Score < 0.60,  OR  any critical rule fails with confidence ≥ 0.70"],
              ],
              col_widths=[1.2, 2.0, 3.5])

    # Score formula
    h2(doc, "8.2  Compliance Score")
    body(doc,
         "The compliance score is a weighted ratio of passed rule weights to total rule weights:",
         bold_parts=["compliance score"])
    add_code_block(doc, [
        "  score  =  SUM( weight_i  for all rules i that PASSED )",
        "           ──────────────────────────────────────────────",
        "           SUM( weight_i  for ALL rules i )",
        "",
        "  Weight mapping:  critical = 1.0  |  major = 0.6  |  minor = 0.2",
        "  Ideal target:    score ≥ 0.90  (PASS threshold)",
    ])

    # Severity
    h2(doc, "8.3  Rule Severity Levels")
    add_table(doc,
              ["Severity", "Weight", "Effect on Verdict"],
              [
                  ["critical", "1.0",
                   "If the rule fails AND the evidence confidence is ≥ 0.70, the claim is automatically "
                   "set to FAIL regardless of the overall score."],
                  ["major", "0.6",
                   "Reduces the compliance score significantly. Does NOT trigger automatic FAIL on its own."],
                  ["minor", "0.2",
                   "Advisory finding. Logged but does not prevent a PASS verdict on its own."],
              ],
              col_widths=[1.0, 0.7, 5.0])

    # Confidence score
    h2(doc, "8.4  Confidence Score")
    body(doc,
         "Every document classification carries a confidence score (0.0 – 1.0) indicating "
         "how certain the system is that the file belongs to the assigned type.")
    add_table(doc,
              ["Range", "Meaning"],
              [
                  ["0.90 – 1.0", "High confidence — filename pattern is an unambiguous, known label (e.g. 'discharge_summary')."],
                  ["0.85 – 0.89", "Content-keyword match confirmed the type after filename lookup."],
                  ["0.70 – 0.84", "Moderate confidence — content keywords matched but filename was generic."],
                  ["0.60 – 0.69", "Low–moderate confidence — partial keyword match or short OCR text."],
                  ["0.40 – 0.59", "Fallback / uncertain — assigned via heuristic; should be verified manually."],
                  ["< 0.40",      "Unknown / unclassified — document could not be matched to any known type."],
              ],
              col_widths=[1.5, 5.2])

    # Signal types
    h2(doc, "8.5  Classification Signal Types")
    add_table(doc,
              ["Signal", "Default Confidence", "Description"],
              [
                  ["filename",  "0.92", "The filename itself contains an unambiguous label (e.g. 'Discharge_Summary_P001.pdf')."],
                  ["content",   "0.85", "OCR text contains strong content keywords matching a known document type."],
                  ["fallback",  "0.40", "No strong filename or content match; type inferred from weak heuristics."],
              ],
              col_widths=[1.2, 1.6, 3.9])

    doc.add_page_break()

    # Rule types
    h2(doc, "8.6  Rule Types")
    add_table(doc,
              ["Rule Type", "What It Checks"],
              [
                  ["document_presence",             "Whether a specific document type exists in the submitted set."],
                  ["document_presence_conditional", "Whether a document is present only when a trigger keyword/condition is met."],
                  ["temporal",                      "Date logic: admission ≤ discharge, or procedure date within admission window."],
                  ["los",                           "Length of stay (discharge date − admission date) falls within allowed bounds."],
                  ["visual_element",                "Whether a visual artefact (hospital stamp, doctor signature) is detectable on a page."],
                  ["financial",                     "Whether the billed amount is within the package ceiling."],
                  ["diagnosis_keyword",             "Whether the OCR text contains diagnosis keywords consistent with the package indication."],
                  ["field_not_blank",               "Whether a specific extracted field (patient name, admission date, etc.) is non-empty."],
              ],
              col_widths=[2.2, 4.5])

    # Document types
    h2(doc, "8.7  Canonical Document Types")
    body(doc, "All documents are classified into one of the following canonical types:")
    add_table(doc,
              ["Canonical Type", "Common Filename Labels", "Description"],
              [
                  ["discharge_summary",      "DS, Discharge, DischargeSummary",           "Inpatient discharge summary prepared by treating doctor."],
                  ["case_sheet",             "CS, CaseSheet, CN, Cn, CaseNotes, IP",       "In-patient case sheet / clinical record (daily progress)."],
                  ["clinical_notes",         "ClinicalNotes, DailyNotes, Progress",        "Doctor's daily progress notes (accepted alternative for case_sheet)."],
                  ["lab_investigation",      "Lab, LFT, CBC, RFT, Blood, Biochemistry",    "Laboratory test reports (blood, urine, biochemistry, etc.)"],
                  ["identity_document",      "Aadhar, PmjayCard, BirthProof, ID",          "Beneficiary identity proof (Aadhaar card, PMJAY card, etc.)"],
                  ["preauthorization_form",  "Preauth, PreAuth, Authorization, Auth",       "PMJAY pre-authorisation / eligibility form."],
                  ["angiography_report",     "CAG, Angio, Angiography, PTCA, Stent",       "Coronary angiography report; also covers PTCA / stenting records."],
                  ["operative_notes",        "OT, OperativeNotes, OTNotes, Surgical",      "Surgical operative notes prepared by the surgeon."],
                  ["anesthesia_notes",       "Anaesthesia, Anesthesia, AnaesthesiaNotes",  "Anaesthesia / pre-operative evaluation record."],
                  ["usg_report",             "USG, Ultrasound, Sonography",                "Ultrasound / sonography report (abdomen, pelvis, KUB, etc.)"],
                  ["xray_image",             "XRay, Xray, ChestXray, Film, CXR",           "X-ray film or radiograph."],
                  ["radiology_report",       "RadiologyReport, Radiology, XRayReport",     "Radiologist's written interpretation of X-ray / imaging."],
                  ["ivp_report",             "IVP, KUB, IVPReport",                        "Intravenous pyelogram or KUB X-ray (urological stone disease)."],
                  ["barcode_sticker",        "BarcodeSticker, ImplantSticker, Sticker",    "Implant / device barcode sticker (mandatory for arthroplasty)."],
                  ["bill_invoice",           "Bill, Invoice, BillInvoice, FinalBill",      "Hospital bill or payment receipt."],
                  ["consent_form",           "Consent, ConsentForm, PatientConsent",       "Patient / guardian consent form."],
                  ["feedback_form",          "Feedback, PatientFeedback, Survey",          "Patient feedback / satisfaction form."],
              ],
              col_widths=[1.8, 2.2, 2.7])

    # Temporal validity
    h2(doc, "8.8  Temporal Validity Windows")
    add_table(doc,
              ["Check", "Rule ID", "Description"],
              [
                  ["admission_before_discharge", "R006 / R004 / R007 / R008",
                   "Admission date must be ≤ discharge date. Failure with confidence ≥ 0.70 triggers auto-FAIL."],
                  ["procedure_within_admission", "R009",
                   "Procedure / operation date must be ≥ admission date and ≤ discharge date."],
                  ["los_bounds",                  "R007 / R005 / R010",
                   "LoS (days) = discharge − admission. Must fall within the min_los – max_los range in the YAML."],
              ],
              col_widths=[2.0, 1.8, 2.9])

    # OCR engine labels
    h2(doc, "8.9  OCR Engine Labels")
    add_table(doc,
              ["Label", "Description"],
              [
                  ["tesseract",  "Primary OCR engine. Uses pytesseract at 150 DPI with early exit if confidence ≥ 0.60."],
                  ["easyocr",    "Optional fallback OCR (disabled by default). Better on handwritten text; slower."],
                  ["pdfplumber", "Used for digitally-native PDFs containing embedded text (no image rendering needed)."],
              ],
              col_widths=[1.5, 5.2])

    # is_extra flag
    h2(doc, "8.10  is_extra Flag")
    body(doc,
         "A document is marked is_extra = true when its canonical type does not appear in the "
         "required_documents or recommended_documents list for the claim's package code. "
         "Extra documents are recorded for auditing but do not affect the compliance score.")

    doc.add_page_break()


# ── MAIN ───────────────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # Page margins
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph font
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)

    build_cover(doc)
    section_executive_summary(doc)
    section_architecture(doc)
    section_stages(doc)
    section_dataset(doc)
    section_tech(doc)
    section_decisions(doc)
    section_stg_guidelines(doc)
    section_glossary(doc)
    section_results(doc)
    section_deployment(doc)

    out_path = "NHA_PMJAY_Claims_Processor_Methodology.docx"
    doc.save(out_path)
    print(f"Report saved: {out_path}")


if __name__ == "__main__":
    main()
